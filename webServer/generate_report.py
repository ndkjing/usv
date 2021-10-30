# encoding: utf-8
"""
@version: 1.0
@author: ndk
@file: generate_report
@time: 2021/10/30
自动生成报告
1 使用数据绘制图片
2 使用图片插入到word中生成报告
"""
import os
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

from docx.enum.text import WD_ALIGN_PARAGRAPH
import re  # 正则表达式

plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号


def generate_docx():
    path = r'F:\pythonProject\usv\statics\imgs'
    Filelist = []
    File_name_lst = []
    for home, dirs, files in os.walk(path):
        for filename in files:
            # 文件名列表，包含完整路径
            Filelist.append(os.path.join(home, filename))
            # # 文件名列表，只包含文件名
            # re.findall(r'\d',str1)
            File_name_lst.append(filename)

    print(Filelist)
    print(File_name_lst)

    document = Document()  # 实例化Document
    document.add_heading('行星轮自动化统计数据', 0)
    for i in range(len(File_name_lst)):
        # 添加图片
        document.add_picture(Filelist[i], width=Inches(5))
        # -----以下代码用来将图片居中----------#
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置

        document.styles['Normal'].font.name = u'黑体'
        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fig_txt = re.findall(r"[\u4e00-\u9fa5]", File_name_lst[i])  # 采用正则表达式，提取文件名中的中文字符。
        run = p.add_run(fig_txt)

    document.add_page_break()

    document.save('result.docx')


def draw_img():
    plt.figure(figsize=(20, 10), dpi=100)
    game = ['1-G1', '1-G2', '1-G3', '1-G4', '1-G5', '2-G1', '2-G2', '2-G3', '2-G4', '2-G5', '3-G1', '3-G2', '3-G3',
            '3-G4', '3-G5', '总决赛-G1', '总决赛-G2', '总决赛-G3', '总决赛-G4', '总决赛-G5', '总决赛-G6']
    scores = [23, 10, 38, 30, 36, 20, 28, 36, 16, 29, 15, 26, 30, 26, 38, 34, 33, 25, 28, 40, 28]
    plt.plot(game, scores, c='red')
    plt.scatter(game, scores, c='red')
    y_ticks = range(50)
    plt.yticks(y_ticks[::5])
    plt.grid(True, linestyle='--', alpha=0.5)
    plt.xlabel("赛程", fontdict={'size': 16})
    plt.ylabel("得分", fontdict={'size': 16})
    plt.title("NBA2020季后赛詹姆斯得分", fontdict={'size': 20})
    plt.savefig(r'F:\pythonProject\usv\statics\imgs\\' + '测试.png')
    plt.show()


if __name__ == '__main__':
    draw_img()
    generate_docx()
